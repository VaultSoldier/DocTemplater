import os
import logging
import random
import sys
import tempfile
import math
from typing import Final, Iterable, List, Optional
from docx.enum.text import WD_BREAK
from docxtpl import DocxTemplate, RichText
from docx import Document
from docxcompose.composer import Composer
from app_logic.processing.data import SqliteData, get_resource_path_temp
from app_logic.types import QuestionType


class DocxProcessingError(Exception):
    """Базовое исключение."""

    pass


class NoQuestionsError(DocxProcessingError):
    pass


class InvalidNumberError(DocxProcessingError):
    pass


class UnknownTicketTypeError(DocxProcessingError):
    pass


class Processing:
    def __init__(self) -> None:
        self.PATH_BASE_DOC: Final[str] = get_resource_path_temp(
            "assets/templates/base.docx"
        )
        self.sql = SqliteData()
        self.is_windows = sys.platform.startswith("win")

        self.practical_questions: list[str] = []
        self.practical_questions_count: int = 0

        self.theoretical_questions: list[str] = []
        self.theoretical_questions_count: int = 0

    def questions_import(self):
        self.practical_questions: list[str] = self.sql.read_questions_list(
            QuestionType.PRACTICAL
        )
        self.practical_questions_count: int = len(self.practical_questions)

        self.theoretical_questions: list[str] = self.sql.read_questions_list(
            QuestionType.THEORETICAL
        )
        self.theoretical_questions_count: int = len(self.theoretical_questions)

    def get_list_safe(
        self, items_list: list, index: int, fallback: Optional[bool] = False
    ) -> str:
        if not items_list:
            return ""

        try:
            item = items_list[index]
            return str(item)
        except IndexError:
            if fallback:
                item = str(random.choice(items_list))
                return item

            return ""

    def get_dict_safe(self, mapping: dict, index: int) -> str:
        try:
            value = mapping[index]
        except KeyError:
            value = ""
        return value

    def process_docx(
        self,
        save_to: str,
        subject: str,
        spec: str,
        cmk: str,
        tutor: str,
        date: list,
        tickets_count: int | None,
        qualify_status: bool | None,
        tickets_count_type: str,
        theoretical_rnd_type: str,
        practical_rnd_type: str,
    ) -> (
        None
        | tuple[
            tempfile._TemporaryFileWrapper,
            List[tempfile._TemporaryFileWrapper],
        ]
    ):
        logging.info(
            f"subject: {subject}\nspec: {spec}\ncmk: {cmk}\ntutor: {tutor}\ndate: {date}\n"
        )

        #  INFO: ОБНОВЛЕНИЕ ВОПРОСОВ
        self.questions_import()

        match tickets_count_type:
            case "Manual" if tickets_count is None or tickets_count <= 0:
                raise InvalidNumberError("Неверное количество билетов")
            case "Practical" if self.practical_questions_count <= 0:
                raise NoQuestionsError("Нету практических вопросов")
            case "Theoretical" if self.theoretical_questions_count <= 0:
                raise NoQuestionsError("Нету теоретических вопросов")

            case "Manual" if tickets_count is not None:
                tickets = range(tickets_count)
            case "Practical":
                tickets = range(self.practical_questions_count)
            case "Theoretical":
                tickets = range(self.theoretical_questions_count)

            case _:
                raise UnknownTicketTypeError(
                    f"Неизвестный тип билетов: {tickets_count_type}"
                )

        qualify = " (квалификационный)" if qualify_status else ""
        day = f"{int(date[2]):02}" or "__"  # add "0" to single num (1 = 01, 10 = 10)
        month = date[1] or ""
        year = date[0] or ""

        rt_day = RichText()
        rt_month = RichText()
        rt_day.add(day, underline=True)

        total_width = 16
        padding = total_width - len(str(month))
        left_pad = " " * (padding // 2)
        right_pad = " " * math.ceil(padding / 2)

        rt_month.add(left_pad)
        rt_month.add(month, underline="thick", bold=True)  # type: ignore[assignment]
        rt_month.add(right_pad)

        context = {
            "qualify": qualify,
            "subject": subject,
            "spec": spec,
            "cmk": cmk,
            "tutor": tutor,
            "day": rt_day,
            "month": rt_month,
            "year": year,
        }
        tpl = DocxTemplate(self.PATH_BASE_DOC)

        if len(tickets) == 1:
            tickets_count_num: int = tickets[0]
            question_one = self.get_selected_questions(
                QuestionType.PRACTICAL, practical_rnd_type, tickets_count_num
            )
            question_two = self.get_selected_questions(
                QuestionType.THEORETICAL, theoretical_rnd_type, tickets_count_num
            )
            context_extend = {
                "ticket_num": str(tickets_count_num + 1),
                "question_one": question_one,
                "question_two": question_two,
            }
            context.update(context_extend)
            tpl.render(context)
            tpl.save(save_to)
            return

        context_extend = {
            "ticket_num": "{{ticket_num}}",
            "question_one": "{{question_one}}",
            "question_two": "{{question_two}}",
        }
        context.update(context_extend)

        tmp_base_docx_file = tempfile.NamedTemporaryFile(
            prefix="tmp_base_",
            suffix=".docx",
            delete=not self.is_windows,
        )

        tpl.render(context)
        tpl.save(tmp_base_docx_file.name)
        tmp_docx_files = self.replace_questions(
            status_rnd_practical=practical_rnd_type,
            status_rnd_theoretical=theoretical_rnd_type,
            tickets=tickets,
            tpl_template_file=tmp_base_docx_file,
        )

        self.docx_merge(tmp_docx_files, save_to)
        return tmp_base_docx_file, tmp_docx_files

    def get_selected_questions(
        self,
        question_type: QuestionType,
        status_rnd: str,
        question_index: int,
    ) -> str:
        if question_type == QuestionType.PRACTICAL:
            questions_list = self.practical_questions
        elif question_type == QuestionType.THEORETICAL:
            questions_list = self.theoretical_questions

        if questions_list is None:
            return ""

        match status_rnd:
            case "fallback":
                question = self.get_list_safe(questions_list, question_index, True)
            case "always":
                question = random.choice(questions_list)
            case "none":
                question = self.get_list_safe(questions_list, question_index)
            case _:
                question = ""
        return str(question)

    def replace_questions(
        self,
        status_rnd_practical: str,
        status_rnd_theoretical: str,
        tickets: range,
        tpl_template_file: tempfile._TemporaryFileWrapper,
    ) -> list[tempfile._TemporaryFileWrapper]:
        """Replace questions"""

        tpl = DocxTemplate(tpl_template_file.name)
        tmpfiles = []

        for i in tickets:
            tmpfile = tempfile.NamedTemporaryFile(
                prefix=f"tmp_{i}", suffix=".docx", delete=not self.is_windows
            )
            question_one = self.get_selected_questions(
                QuestionType.PRACTICAL, status_rnd_practical, i
            )
            question_two = self.get_selected_questions(
                QuestionType.THEORETICAL, status_rnd_theoretical, i
            )

            context = {
                "ticket_num": f"{i+1}",
                "question_one": question_one,
                "question_two": question_two,
            }
            tpl.render(context)
            tpl.save(tmpfile.name)
            tmpfiles.append(tmpfile)
        return tmpfiles

    def docx_merge(
        self, paths: list[tempfile._TemporaryFileWrapper], save_to: str
    ) -> None:
        """Merge temporary files"""
        master = Document(paths[0])
        composer = Composer(master)

        # master file page break
        master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        for idx, tmp in enumerate(paths[1:], start=1):
            doc = Document(tmp.name)
            composer.append(doc)
            if idx != len(paths) - 1:
                master.add_page_break()
        composer.save(save_to)

    def clean(
        self,
        path: Optional[tempfile._TemporaryFileWrapper] = None,
        paths: Optional[Iterable[tempfile._TemporaryFileWrapper]] = None,
    ):
        if path is None and paths is None:
            return

        all_paths = []
        if path is not None:
            all_paths.append(path)
        if paths is not None:
            all_paths.extend(paths)

        # clean temporary files
        for file in all_paths:
            try:
                file_path = file.name
                # Deletes files on UNIX
                # And makes sure it closes before deleting on non UNIX
                if not file.closed:
                    file.close()
                # Deletes files on non UNIX systems
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                logging.error(f"Failed to delete {file.name}: {e}")
